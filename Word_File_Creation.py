from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
import json
import io
import re
from typing import Union, Dict, Any
import ast

#========================== ALL METHOD IN THE JSON FILE IS =====
'''
#def generate_resume(json_file_path, output_file='resume.docx'): TO CRAETE THE WORD FILW WHIC WILL ACCEPT THE JSON FILE
#def convert_to_json_string(json_response):
'''
#================================================================
import re
import os

def extract_and_save_json_to_file(final_response: str):
    """
    Extracts the JSON block from a raw LLM response string and saves it as a .json file
    in the specified directory, using the value of the "name" key as the filename.

    Args:
        final_response (str): Raw string like 'JSON = {...} Missing_Technical_Skills = [...]'

    Returns:
        str: File path if saved successfully, else None
    """
    try:
        # Extract the JSON block using regex
        match = re.search(r'JSON\s*=\s*(\{.*?\})(?=\s*Missing_|$)', final_response, re.DOTALL)
        if not match:
            raise ValueError("JSON block not found in response")

        json_text = match.group(1)

        # Extract the name value from the JSON string without converting to dict
        name_match = re.search(r'"name"\s*:\s*"([^"]+)"', json_text)
        if not name_match:
            raise ValueError("Name key not found in JSON")

        name_value = name_match.group(1)

        # Define the file path
        output_dir = "/Users/syamvemuri/Downloads/Sayyam/ai/Spartex_AI/JSON_FILES"
        os.makedirs(output_dir, exist_ok=True)  # Ensure directory exists
        file_path = os.path.join(output_dir, f"{name_value}.json")

        # Save to file
        with open(file_path, "w") as file:
            file.write(json_text)

        print(f"JSON saved to {file_path}")
        return file_path

    except Exception as e:
        print(f"Error extracting or saving JSON: {e}")
        return None

#================= START METHOD TO  GENARATE RESUME 


def generate_resume_from_json(json_file_path, output_dir,font_style):
    # Load JSON data
    with open(json_file_path, 'r') as file:
        data = json.load(file)

    name = data.get('name', 'resume').replace(" ", "_")
    output_path = os.path.join(output_dir, f"{name}.docx")
        
    # Initialize document
    doc = Document()

    # Set tight margins (0.5 inches)
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Configure base style
    style = doc.styles['Normal']
    font = style.font
    font.name = font_style
    font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # Helper functions
    def add_centered_paragraph(text, bold=False, size=11, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.bold = True
        return p

    def add_justified_paragraph(text=None, bold=False, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if text:
            run = p.add_run(text)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)
        return p

    def add_bullet_points(items):
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(item)
            run.font.size = Pt(11)

    # --- Personal Information ---
    add_centered_paragraph(data['name'], bold=True, size=16)
    add_centered_paragraph(data['title'], size=12)
    contact = data['contact']
    contact_parts = []
    if 'email' in contact:
        contact_parts.append(contact['email'])
    if 'phone' in contact:
        contact_parts.append(contact['phone'])
    if 'linkedin' in contact:
        contact_parts.append('LinkedIn')
    if contact.get('portfolio'):  # This handles None and empty strings
        contact_parts.append('Portfolio')

    contact_text = " | ".join(contact_parts)
    add_centered_paragraph(contact_text, size=10)

    # --- Professional Summary ---
    if 'professional_summary' in data and data['professional_summary']:
        add_section_heading("Professional Summary")
        add_bullet_points(data['professional_summary'])

    # --- Technical Skills ---
    if 'technical_skills' in data and data['technical_skills']:
        add_section_heading("Technical Skills")
        for category, skills in data['technical_skills'].items():
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(f"{category}: ")
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            run = p.add_run(", ".join(skills))
            run.font.size = Pt(11)
            run.bold = False

    # --- Professional Experience ---
    if 'experience' in data and data['experience']:
        add_section_heading("Professional Experience")
        for job in data['experience']:
            header_table = doc.add_table(rows=1, cols=2)
            header_table.autofit = False
            header_table.columns[0].width = Inches(4.5)
            header_table.columns[1].width = Inches(1.5)

            left_cell = header_table.cell(0, 0)
            role_run = left_cell.paragraphs[0].add_run(f"Role: {job['role']}")
            role_run.bold = True
            role_run.font.color.rgb = RGBColor(0, 0, 0)

            client_info = f"Client: {job['client']}"
            if 'location' in job and job['location']:
                client_info += f" – {job['location']}"
            client_para = left_cell.add_paragraph()
            client_run = client_para.add_run(client_info)
            client_run.bold = True
            client_run.font.color.rgb = RGBColor(0, 0, 0)

            right_cell = header_table.cell(0, 1)
            right_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            duration_run = right_cell.paragraphs[0].add_run(job['duration'])
            duration_run.bold = True
            duration_run.font.color.rgb = RGBColor(0, 0, 0)

            if 'responsibilities' in job and job['responsibilities']:
                if job.get('client') == "Cipla Pharmaceuticals":
                    p = add_justified_paragraph("Responsibilities: To modernize Cipla's healthcare data infrastructure", bold=True)
                else:
                    p = add_justified_paragraph("Responsibilities:", bold=True)
                add_bullet_points(job['responsibilities'])

            if 'environment' in job and job['environment']:
                p = add_justified_paragraph("Environment:", bold=True)
                p = add_justified_paragraph(", ".join(job['environment']))

    # --- Education ---
    if 'education' in data:
        add_section_heading("Education")
        educations = data['education'] if isinstance(data['education'], list) else [data['education']]

        for edu in educations:
            p = doc.add_paragraph()
            if 'degree' in edu:
                degree_run = p.add_run(f"{edu['degree']} ")
                degree_run.bold = True
                degree_run.font.color.rgb = RGBColor(0, 0, 0)

                remaining_parts = []
                if 'field' in edu:
                    remaining_parts.append(f"in {edu['field']}")
                if 'concentration' in edu and edu['concentration']:
                    remaining_parts.append(f"({edu['concentration']})")
                if 'institution' in edu:
                    remaining_parts.append(f"| {edu['institution']}")
                if 'year' in edu and edu['year'] and str(edu['year']).strip().upper() != 'YYYY':
                    remaining_parts.append(f"({edu['year']})")

                remaining_text = " ".join(remaining_parts)
                p.add_run(remaining_text)

    # --- Certifications ---
    if 'certifications' in data and data['certifications']:
        add_section_heading("Certifications")
        add_bullet_points(data['certifications'])

    os.makedirs(output_dir, exist_ok=True)
    doc.save(output_path)
    print(f" Resume saved at: {output_path}")
    return output_path
    

#json_file =  ast.literal_eval(json_file)
#File_path = '/Users/syamvemuri/Downloads/Sayyam/ai/Spartex_AI/Resumes'
#final_resume_file = generate_resume(json_file,file_path)
#print(f' the final resume file is ====> {final_resume_file}')
# Word_File_Creation.py